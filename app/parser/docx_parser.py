"""
Извлечение структурированного контента из DOCX без исполнения макросов.

Порядок блоков сохраняется как в документе: заголовки, абзацы, inline-
иллюстрации на своих местах. Загруженные отдельно картинки привязываются
в app.nlp.image_matcher.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image

from app.layout.image_roles import classify_image


HEADING_RE = re.compile(r"^Heading\s*(\d)$", re.IGNORECASE)
PAGE_BREAK_MARKER = re.compile(
    r"^\s*(\[PAGE[_\s]?BREAK\]|\[НОВАЯ[_\s]?СТРАНИЦА\]|---+\s*PAGE\s*---+|=+\s*PAGE\s*=+)\s*$",
    re.IGNORECASE,
)


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class Block:
    kind: str  # heading | paragraph | list_item | image | caption | page_break | table | footnote_block
    level: int = 0
    runs: list[Run] = field(default_factory=list)
    image_index: int | None = None
    caption: str = ""
    image_role: str = "photo"  # photo | banner | logo | ad
    width_mm: float | None = None
    height_mm: float | None = None
    slot_index: int | None = None
    table_rows: list[list[list[Run]]] = field(default_factory=list)
    footnote_ids: list[int] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.runs)


@dataclass
class ExtractedImage:
    path: Path
    width_px: int
    height_px: int
    dpi: tuple[int, int]
    source: str = "inline"  # inline | upload | marker | stock
    original_name: str = ""
    role: str = "photo"
    width_mm: float | None = None
    height_mm: float | None = None


@dataclass
class ParsedDocument:
    blocks: list[Block]
    images: list[ExtractedImage]
    footnotes: dict[int, list[Run]] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        parts = []
        for b in self.blocks:
            if b.kind == "image":
                continue
            if b.kind == "table":
                for row in b.table_rows:
                    parts.append(" | ".join("".join(r.text for r in cell) for cell in row))
            else:
                parts.append(b.text)
        return "\n".join(parts)

    @property
    def word_count(self) -> int:
        return len(re.findall(r"\w+", self.full_text, flags=re.UNICODE))


class DocxParseError(Exception):
    pass


def _paragraph_has_page_break(paragraph: Paragraph) -> bool:
    for br in paragraph._element.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _iter_body_items(document: DocumentType):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _style_to_block_kind(paragraph: Paragraph) -> tuple[str, int]:
    style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
    m = HEADING_RE.match(style_name)
    if m:
        return "heading", int(m.group(1))
    if style_name.lower() in ("list paragraph", "list bullet", "list number"):
        return "list_item", 0
    text = paragraph.text.strip()
    if text and len(text) < 90 and not text.endswith((".", ",", ";", ":")):
        runs = paragraph.runs
        if runs and all(r.bold for r in runs if r.text.strip()):
            return "heading", 3
    return "paragraph", 0


def _extract_runs(paragraph: Paragraph) -> list[Run]:
    out: list[Run] = []
    footnote_ids: list[int] = []
    for r in paragraph.runs:
        for fn in r._element.findall(".//" + qn("w:footnoteReference")):
            try:
                fid = int(fn.get(qn("w:id"), "0"))
                footnote_ids.append(fid)
                out.append(Run(text=str(len(footnote_ids)), bold=False, italic=False))
            except (TypeError, ValueError):
                pass
        if not r.text:
            continue
        out.append(Run(text=r.text, bold=bool(r.bold), italic=bool(r.italic)))
    return out


def _extract_runs_from_cell(cell: _Cell) -> list[Run]:
    parts: list[Run] = []
    for p in cell.paragraphs:
        parts.extend(_extract_runs(p) or ([Run(text=p.text)] if p.text.strip() else []))
    return parts


def _parse_table_rows(table: Table) -> list[list[list[Run]]]:
    rows: list[list[list[Run]]] = []
    for row in table.rows:
        cells: list[list[Run]] = []
        for cell in row.cells:
            cell_runs = _extract_runs_from_cell(cell)
            if not cell_runs and cell.text.strip():
                cell_runs = [Run(text=cell.text.strip())]
            cells.append(cell_runs)
        if any(c for c in cells):
            rows.append(cells)
    return rows


def _blip_embed_ids(paragraph: Paragraph) -> list[str]:
    ids = []
    for blip in paragraph._element.findall(".//" + qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if embed:
            ids.append(embed)
    return ids


def _save_image_blob(blob: bytes, images_out_dir: Path, images: list[ExtractedImage]) -> int:
    images_out_dir.mkdir(parents=True, exist_ok=True)
    img = Image.open(io.BytesIO(blob))
    img.load()
    ext = (img.format or "PNG").lower()
    ext = "jpg" if ext == "jpeg" else ext
    fname = f"inline_{len(images) + 1:02d}.{ext}"
    out_path = images_out_dir / fname
    img.save(out_path)
    dpi = img.info.get("dpi", (96, 96))
    images.append(ExtractedImage(
        path=out_path,
        width_px=img.width,
        height_px=img.height,
        dpi=(int(dpi[0]), int(dpi[1])),
        source="inline",
        original_name=fname,
        role=classify_image(out_path, fname),
    ))
    return len(images) - 1


def _image_index_from_embed(doc: DocumentType, embed_id: str,
                             images_out_dir: Path, images: list[ExtractedImage],
                             cache: dict[str, int]) -> int | None:
    if embed_id in cache:
        return cache[embed_id]
    try:
        part = doc.part.related_parts[embed_id]
        idx = _save_image_blob(part.blob, images_out_dir, images)
        cache[embed_id] = idx
        return idx
    except Exception:
        return None


def parse_docx(path: Path, images_out_dir: Path) -> ParsedDocument:
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise DocxParseError(f"Не удалось открыть DOCX-файл: {exc}") from exc

    blocks: list[Block] = []
    images: list[ExtractedImage] = []
    embed_cache: dict[str, int] = {}
    from app.parser.footnotes import load_footnotes_blob
    footnotes = load_footnotes_blob(doc)
    collected_fn_ids: list[int] = []

    for item in _iter_body_items(doc):
        if isinstance(item, Paragraph):
            if _paragraph_has_page_break(item):
                blocks.append(Block(kind="page_break", level=0, runs=[]))

            embed_ids = _blip_embed_ids(item)
            text = item.text.strip()

            if text:
                if PAGE_BREAK_MARKER.match(text):
                    blocks.append(Block(kind="page_break", level=0, runs=[]))
                else:
                    kind, level = _style_to_block_kind(item)
                    runs = _extract_runs(item) or [Run(text=item.text)]
                    fn_ids = []
                    for r in item.runs:
                        for fn in r._element.findall(".//" + qn("w:footnoteReference")):
                            try:
                                fn_ids.append(int(fn.get(qn("w:id"), "0")))
                            except (TypeError, ValueError):
                                pass
                    collected_fn_ids.extend(fn_ids)
                    blocks.append(Block(
                        kind=kind, level=level, runs=runs, footnote_ids=fn_ids,
                    ))

            for embed_id in embed_ids:
                img_idx = _image_index_from_embed(doc, embed_id, images_out_dir, images, embed_cache)
                if img_idx is not None:
                    role = images[img_idx].role
                    blocks.append(Block(
                        kind="image", level=0, runs=[],
                        image_index=img_idx, image_role=role,
                    ))

            if not text and not embed_ids:
                continue

        elif isinstance(item, Table):
            table_rows = _parse_table_rows(item)
            if table_rows:
                blocks.append(Block(kind="table", level=0, table_rows=table_rows))

    if collected_fn_ids and footnotes:
        seen: set[int] = set()
        fn_blocks: list[Block] = []
        for fid in collected_fn_ids:
            if fid in seen or fid not in footnotes:
                continue
            seen.add(fid)
            marker = str(len(seen))
            fn_blocks.append(Block(
                kind="footnote_block",
                runs=[Run(text=f"{marker}. ", bold=False), *footnotes[fid]],
            ))
        if fn_blocks:
            blocks.append(Block(kind="heading", level=3, runs=[Run(text="Примечания", bold=True)]))
            blocks.extend(fn_blocks)

    if not blocks:
        raise DocxParseError("В документе не найдено текстового содержимого.")

    return ParsedDocument(blocks=blocks, images=images, footnotes=footnotes)
